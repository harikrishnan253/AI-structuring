"""
STAGE 1: Document Ingestion
- DOCX Parser (python-docx)
- Structure Extractor
- Document Serializer
- Context Zone Detection (Front Matter, Body, Back Matter, Table, Box)

Extracts all paragraphs with formatting metadata and assigns unique IDs.
Detects document context zones to provide better classification context.

CONTEXT ZONES:
- FRONT_MATTER: Metadata, author info → PMI only
- BODY: Main chapter content → Full style range
- BACK_MATTER: References, index → REF-N, EOC-*, IX-*
- TABLE: Word table content → T, T2, T4, TBL-*, TFN, TSN
- BOX: Pedagogical boxes → NBX-*, BX1-*, BX2-*, BX3-*
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Optional
import re
import logging

logger = logging.getLogger(__name__)


# =============================================================================
# ZONE DETECTION PATTERNS
# =============================================================================

# Patterns that indicate FRONT MATTER (metadata, author info, etc.)
FRONT_MATTER_PATTERNS = [
    r'^book\s*title:',
    r'^chapter\s*#:',
    r'^chapter\s*title:',
    r'^corresponding\s*author:',
    r'^orcid\s*identifier:',
    r'^<metadata>',
    r'^key\s*words:',
    r'^abstract:',
    r'^section:$',
    r'^phone:\s*\(\d{3}\)',
    r'@.*\.(org|com|edu|net)$',  # Email patterns
    r'^\d{5}$',  # ZIP codes
]

# Patterns that indicate BACK MATTER (references, index, etc.)
BACK_MATTER_PATTERNS = [
    r'^<ref>references',
    r'^references\s*$',
    r'^bibliography\s*$',
    r'^index\s*$',
    r'^appendix\s+[a-z]',
    r'^glossary\s*$',
    r'^suggested\s+reading',
    r'^further\s+reading',
    r'^answer\s+key',
]

# Patterns that indicate CHAPTER OPENER (CN/CT - start of front matter)
CHAPTER_OPENER_PATTERNS = [
    r'^<cn>chapter\s+\d+',
    r'^chapter\s+\d+\s*$',
    r'^chapter\s+\d+[:\s]',
    r'^\d+\s*$',  # Just a number (chapter number)
]

# Patterns that indicate BODY START (first H1 heading - main content begins)
BODY_START_PATTERNS = [
    r'^<h1>',  # Explicit H1 marker
    r'^<h1\s',  # H1 with attributes
    r'^\s*introduction\s*$',  # Common first H1
    r'^\s*overview\s*$',
    r'^\s*background\s*$',
    r'^\s*getting\s+started\s*$',
]

# Patterns that indicate METADATA (pure PMI - before chapter opener)
METADATA_PATTERNS = [
    r'^book\s*title:',
    r'^chapter\s*#:',
    r'^chapter\s*title:',
    r'^corresponding\s*author:',
    r'^orcid\s*identifier:',
    r'^<metadata>',
    r'^key\s*words:',
    r'^abstract:',
    r'^section:\s*$',
    r'^phone:\s*\(\d{3}\)',
    r'@.*\.(org|com|edu|net)$',  # Email patterns
    r'^\d{5}$',  # ZIP codes
]

# Patterns that indicate BOX START
BOX_START_PATTERNS = [
    r'^<note>',
    r'^<clinical\s*pearl>',
    r'^<red\s*flag>',
    r'^<box>',
    r'^<tip>',
    r'^<example>',
    r'^<warning>',
    r'^<alert>',
    r'^<case\s*study>',
    r'^<reflection>',
    r'^<discussion>',
    r'^<practice>',
    r'^<key\s*point>',
    r'^<important>',
    r'^<remember>',
    r'^<unnumbered\s*box>',
]

# Patterns that indicate BOX END
BOX_END_PATTERNS = [
    r'^</note>',
    r'^</clinical\s*pearl>',
    r'^</red\s*flag>',
    r'^</box>',
    r'^</tip>',
    r'^</example>',
    r'^</warning>',
    r'^</alert>',
    r'^</case\s*study>',
    r'^</reflection>',
    r'^</discussion>',
    r'^</practice>',
    r'^</key\s*point>',
    r'^</important>',
    r'^</remember>',
    r'^</unnumbered\s*box>',
]

# Box type to style prefix mapping
BOX_TYPE_MAPPING = {
    'note': 'NBX',
    'tip': 'NBX',
    'clinical pearl': 'BX1',
    'example': 'BX1',
    'red flag': 'BX2',
    'warning': 'BX2',
    'alert': 'BX2',
    'reflection': 'BX3',
    'discussion': 'BX3',
    'case study': 'BX4',
    'practice': 'EXER',
    'key point': 'NBX',
    'important': 'NBX',
    'remember': 'NBX',
    'unnumbered box': 'NBX',
    'box': 'NBX',
}

# =============================================================================
# ZONE-SPECIFIC VALID STYLES 
# Extracted from 34 training documents + Official WK Book Template 1.1
# =============================================================================

ZONE_VALID_STYLES = {
    # METADATA: Pure pre-press info before chapter opener
    'METADATA': [
        'PMI',
        # Author/title info that appears before chapter
        'ChapterAuthor', 'ChapterNumber', 'ChapterTitle', 'ChapterTitleFootnote',
        # Special headings for abstract, keywords
        'SP-Heading2', 'SpeacialHeading2', 'SpecialHeading2',
        # Paragraph styles
        'Para-FL', 'ParaFirstLine-Ind',
        # Markers
        'H2', 'Normal',
    ],
    
    # FRONT_MATTER: Chapter opener through objectives (before first H1)
    'FRONT_MATTER': [
        # Chapter opener (Official WK Template)
        'CN', 'CT', 'CST', 'CAU', 'CW', 'CHAP', 'COQ', 'COQA',
        'ChapterNumber', 'ChapterTitle', 'ChapterAuthor',
        'Chapter-Epigraph', 'Chapter-EpigraphSource', 'EPIGRAPH',
        # Part openers (Official WK Template)
        'PART', 'PN', 'PT', 'PST', 'PAU', 'PTXT', 'PTXT-DC',
        'POC', 'POC-FIRST', 'POUT-1', 'POUT-2', 'POUTH1', 'PQUOTE', 'POS',
        # Unit openers (Official WK Template)
        'UNIT', 'UN', 'UT', 'UST', 'UAU', 'UTXT',
        'UOC', 'UOC-FIRST', 'UOUT-1', 'UOUT-2', 'UOUTH1', 'UQUOTE', 'UOS',
        # Section openers (Official WK Template)
        'SECTION', 'SN', 'ST', 'SST', 'SAU', 'STXT', 'STXT-DC',
        'SOC', 'SOC-FIRST', 'SOUT-1', 'SOUT-2', 'SOUTH1', 'SQUOTE', 'SOS',
        'SOUT-NL-FIRST', 'SOUT-NL-MID',
        # Objectives - Bulleted (Official WK Template)
        'OBJ1', 'OBJ-TXT', 'OBJ-TXT-FLUSH',
        'OBJ-BL-FIRST', 'OBJ-BL-MID', 'OBJ-BL-LAST',
        # Objectives - Numbered (Official WK Template)
        'OBJ-NL-FIRST', 'OBJ-NL-MID', 'OBJ-NL-LAST', 'OBJ_NL',
        # Objectives - Unnumbered (Official WK Template)
        'OBJ-UL-FIRST', 'OBJ-UL-MID', 'OBJ-UL-LAST',
        # Chapter outline (Official WK Template)
        'COUT-1', 'COUT-2', 'COUT-3', 'COUT1', 'COUT2',
        'COUT-1-TXT', 'COUT-1-H1', 'COUTH1',
        'COUT-BL', 'COUT-NL-FIRST', 'COUT-NL-MID',
        # Key Terms (Official WK Template)
        'KT1', 'KT-TXT',
        'KT-BL-FIRST', 'KT-BL-MID', 'KT-BL-LAST',
        'KT-NL-FIRST', 'KT-NL-MID', 'KT-NL-LAST',
        'KT-UL-FIRST', 'KT-UL-MID', 'KT-UL-LAST',
        # Compact objectives (Training data)
        'COBJ', 'COBJ_T', 'COBJ_TXL',
        # Learning objectives alternatives (Training data)
        'LearningObj-BulletList1', 'LearningObj-BulletList1_first', 
        'LearningObj-BulletList1_last', 'LearningObj-BulletList1-last',
        'LearningObj-Para-FL', 'LearnObjHeading',
        # Objectives alternatives (Training data)
        'ObjectiveHead', 'ObjectivesHeading', 'Objectives-Para-FL',
        'ObjectiveNumberList', 'ObjectiveNumberList-First', 'ObjectiveNumberList-Last',
        'ObjectivesNumberlist1', 'ObjectivesNumberlist1_first', 'ObjectivesNumberlist1_last',
        # Special intro elements
        'H1A', 'SP-H1', 'TXT-DC', 'TXT-FLUSH', 'TXT', 'KEYNOTE', 'INTRO',
        'H2', 'H3',
        # Case study in intro
        'CS-TTL', 'CS-TXT-FLUSH', 'CS-TXT', 'CS-H1',
        'CS-QUES-TXT', 'CS-ANS-TXT',
        # Domain title (Training data)
        'DOM-TTL',
        # Markers
        'PMI', 'Normal',
    ],
    
    # TABLE: Word table content
    'TABLE': [
        # Basic table cells (Official WK Template)
        'T', 'T-DIR',
        'T1', 'T11', 'T12',
        'T2', 'T2-C', 'T21', 'T22', 'T23',
        'T3', 'T4', 'T5', 'T6',
        # Table bullets (Official WK Template + Training)
        'TBL', 'TBL-FIRST', 'TBL-MID', 'TBL-MID0', 'TBL-LAST', 'TBL-LAST1',
        'TBL2-MID', 'TBL3-MID', 'TBL4-MID',
        # Table numbered lists (Official WK Template)
        'TNL-FIRST', 'TNL-MID',
        # Table unnumbered lists (Official WK Template + Training)
        'TUL', 'TUL-FIRST', 'TUL-MID', 'TUL-LAST',
        # Table footnotes and sources (Official WK Template)
        'TFN', 'TFN1', 'TFN-FIRST', 'TFN-MID', 'TFN-LAST',
        'TFN-BL-FIRST', 'TFN-BL-MID', 'TFN-BL-LAST',
        'TSN',
        # Table math (Official WK Template)
        'TMATH',
        # Unnumbered tables (Official WK Template)
        'UNT', 'UNT-TTL', 'UNT-T1', 'UNT-T2', 'UNT-T3',
        'UNT-BL', 'UNT-BL-MID',
        'UNT-NL-FIRST', 'UNT-NL-MID',
        'UNT-UL', 'UNT-FN',
        # Unnumbered table in box (Official WK Template)
        'UNBX-TT', 'UNBX-T', 'UNBX-T2', 'UNBX-BL', 'UNBX-NL', 'UNBX-UL',
        # Alternative table styles (Training data)
        'TB', 'TB-BulletList1', 'TB-NumberList1', 'TB-AlphaList1',
        'TableBody', 'TableCaption', 'TableCaptions', 'TableColumnHead1',
        'TableFootnote', 'TableList', 'TableNote', 'TableSource',
        'Exhibit-TableBody', 'Exhibit-TableColumnHead1', 'Exhibit-TB-BulletList1',
        'Exhibit-TableFootnote', 'ExhibitTitle',
        # Clinical judgment in tables (Training data)
        'CJC-UL-MID', 'CJC-UNT', 'CJC-UNBX-T', 'CJC-UNBX-T2',
        # Box content in tables
        'NBX1-UNT', 'NBX1-UNT-T2',
    ],
    
    # BOX_NBX: Informational boxes (NOTE, TIP, etc.)
    'BOX_NBX': [
        # Structure (Official WK Template)
        'NBX-TTL', 'NBX-TYPE', 'NBX-TXT', 'NBX-TXT-DC', 'NBX-TXT-FIRST', 'NBX-TXT-FLUSH',
        'NBX-FN', 'NBX-QUO',
        # Headings (Official WK Template)
        'NBX-H1', 'NBX-H2', 'NBX-H3', 'NBX-L1',
        # Bulleted lists (Official WK Template)
        'NBX-BL-FIRST', 'NBX-BL-MID', 'NBX-BL-LAST', 'NBX-BL2-MID',
        # Numbered lists (Official WK Template)
        'NBX-NL-FIRST', 'NBX-NL-MID', 'NBX-NL-LAST',
        # Unnumbered lists (Official WK Template)
        'NBX-UL-FIRST', 'NBX-UL-MID', 'NBX-UL-LAST',
        # Multi-column lists (Official WK Template)
        'NBX-MCUL-FIRST', 'NBX-MCUL-MID', 'NBX-MCUL-LAST',
        # Outline lists (Official WK Template)
        'NBX-OUT1-FIRST', 'NBX-OUT1-MID', 'NBX-OUT2', 'NBX-OUT2-LAST', 'NBX-OUT3',
        # Equations (Official WK Template)
        'NBX-EQ-FIRST', 'NBX-EQ-MID', 'NBX-EQ-LAST', 'NBX-EQ-ONLY',
        # Extracts (Official WK Template)
        'NBX-EXT-ONLY',
        # Dialogue (Training data)
        'NBX-DIA', 'NBX-DIA-FIRST', 'NBX-DIA-MID', 'NBX-DIA-LAST',
        # Table/source in box
        'NBX-UNT', 'NBX-UNT-T2', 'NBX-SRC',
        # NBX1 variants - Edwards template (Training data)
        'NBX1-TTL', 'NBX1-TXT', 'NBX1-TXT-FIRST', 'NBX1-TXT-FLUSH',
        'NBX1-BL-FIRST', 'NBX1-BL-MID', 'NBX1-BL-LAST', 'NBX1-BL2-MID',
        'NBX1-NL-FIRST', 'NBX1-NL-MID', 'NBX1-NL-LAST',
        'NBX1-DIA-FIRST', 'NBX1-DIA-MID', 'NBX1-DIA-LAST',
        'NBX1-UNT', 'NBX1-UNT-T2', 'NBX1-SRC',
        # Box-01 variants - Wheeler template (Training data)
        'Box-01-BoxTitle', 'Box-01-BulletList1', 'Box-01-BulletList1_first', 'Box-01-BulletList1_last',
        'Box-01-NumberList1', 'Box-01-ParaFirstLine-Ind', 'Box-01-Para-FL',
        'Box-01-Head1',
        'Box-01-UN-TableBody', 'Box-01-UN-TableCaption', 'Box-01-UN-TableColumnHead1', 'Box-01-UN-TableFootnote',
        # Markers
        'PMI',
    ],
    
    # BOX_BX1: Clinical/practical boxes (CLINICAL PEARL, EXAMPLE, TIP)
    'BOX_BX1': [
        # Structure (Official WK Template)
        'BX1-TTL', 'BX1-TYPE', 'BX1-TXT', 'BX1-TXT-DC', 'BX1-TXT-FIRST', 'BX1-TXT-FLUSH',
        'BX1-FN', 'BX1-QUO',
        # Headings (Official WK Template)
        'BX1-H1', 'BX1-H2', 'BX1-H3', 'BX1-L1',
        # Bulleted lists (Official WK Template)
        'BX1-BL-FIRST', 'BX1-BL-MID', 'BX1-BL-LAST', 'BX1-BL2-MID',
        # Numbered lists (Official WK Template)
        'BX1-NL-FIRST', 'BX1-NL-MID', 'BX1-NL-LAST',
        # Unnumbered lists (Official WK Template)
        'BX1-UL-FIRST', 'BX1-UL-MID', 'BX1-UL-LAST',
        # Multi-column lists (Official WK Template)
        'BX1-MCUL-FIRST', 'BX1-MCUL-MID', 'BX1-MCUL-LAST',
        # Outline lists (Official WK Template)
        'BX1-OUT1-FIRST', 'BX1-OUT1-MID', 'BX1-OUT2', 'BX1-OUT2-LAST', 'BX1-OUT3',
        # Equations (Official WK Template)
        'BX1-EQ-FIRST', 'BX1-EQ-MID', 'BX1-EQ-LAST', 'BX1-EQ-ONLY',
        # Extracts (Official WK Template)
        'BX1-EXT-ONLY',
        # Questions in box (Training data)
        'BX1-QUES-TXT',
        # Markers
        'PMI',
    ],
    
    # BOX_BX2: Warning boxes (RED FLAG, WARNING, ALERT)
    'BOX_BX2': [
        'BX2-TTL', 'BX2-TYPE', 'BX2-TXT', 'BX2-TXT-FIRST', 'BX2-TXT-FLUSH', 'BX2-TXT-LAST',
        'BX2-H1',
        'BX2-BL-FIRST', 'BX2-BL-MID', 'BX2-BL-LAST',
        'BX2-NL-FIRST', 'BX2-NL-MID', 'BX2-NL-LAST',
        'PMI',
    ],
    
    # BOX_BX3: Reflection/discussion boxes
    'BOX_BX3': [
        'BX3-TTL', 'BX3-TYPE', 'BX3-TXT', 'BX3-TXT-FIRST', 'BX3-TXT-FLUSH',
        'BX3-BL-FIRST', 'BX3-BL-MID', 'BX3-BL-LAST',
        'BX3-NL-FIRST', 'BX3-NL-MID', 'BX3-NL-LAST',
        'BX3_BL', 'BX3_BLF', 'BX3_BLL',
        'PMI',
    ],
    
    # BOX_BX4: Procedure/Case study boxes
    'BOX_BX4': [
        'BX4-TTL', 'BX4-TYPE', 'BX4-TXT', 'BX4-TXT-FIRST', 'BX4-TXT-FLUSH',
        'BX4-H1', 'BX4-H2',
        'BX4-BL-MID', 'BX4-BL2-MID',
        'BX4-NL-MID', 'BX4-LL2-MID',
        # Case study styles
        'CS-H1', 'CS-TTL', 'CS-TXT', 'CS-TXT-FLUSH',
        'CS-QUES-TXT', 'CS-ANS-TXT',
        'CaseStudy-UL-FL1', 'CaseStudy-Dialogue', 'CaseStudy-Heading1',
        'CaseStudy-ParaFirstLine-Ind', 'CaseStudy-BulletList1',
        'PMI',
    ],
    
    # BOX_BX6: Resource boxes
    'BOX_BX6': [
        'BX6-TTL', 'BX6-TYPE', 'BX6-TXT', 'BX6-TXT-FIRST',
        'BX6-BL-MID',
        'PMI',
    ],
    
    # BOX_BX7: Case study boxes
    'BOX_BX7': [
        'BX7-TTL', 'BX7-TYPE', 'BX7-TXT', 'BX7-TXT-FIRST',
        'BX7-BL-FIRST', 'BX7-BL-LAST',
        'BX7-NL-MID',
        'PMI',
    ],
    
    # BOX_BX15: Special boxes
    'BOX_BX15': [
        'BX15-TTL', 'BX15-TYPE', 'BX15-TXT', 'BX15-TXT-FIRST',
        'BX15-H1',
        'PMI',
    ],
    
    # BOX_BX16: Special boxes with unnumbered tables
    'BOX_BX16': [
        'BX16-TTL', 'BX16-TYPE',
        'BX16-UNT', 'BX16-UNT2', 'BX16-UNT-BL-MID',
        'PMI',
    ],
    
    # BACK_MATTER: References, figures, end-of-chapter, appendix, glossary, index
    'BACK_MATTER': [
        # Reference headings (Official WK Template + Training)
        'REF-H1', 'REF-H2', 'REFH1', 'REFH2', 'REFH2a',
        'H1-REF', 'ReferencesHeading1',
        # Reference entries (Official WK Template + Training)
        'REF-N', 'REF-N-FIRST', 'REF-N0', 'REF', 'REF-U',
        'Reference-Alphabetical', 'ReferenceAlphabetical', 'Reference-Numbered',
        # Suggested readings (Official WK Template)
        'SR', 'SRH1', 'SRH2',
        # Bibliography (Official WK Template)
        'BIB', 'BIBH1', 'BIBH2',
        # Acknowledgments (Official WK Template)
        'ACK1', 'ACKTXT',
        # Web links (Official WK Template)
        'WEBTXT', 'WL1',
        # Figure elements
        'FIG-LEG', 'FIG-CRED', 'FIG-SRC', 'UNFIG', 'UNFIG-LEG', 'UNFIG-SRC',
        'FigureLegend', 'FigureCaption', 'FigureSource', 'FG-CAP',
        # Table elements in back matter
        'T1', 'TFN', 'TSN',
        # Exhibit elements
        'ExhibitTitle', 'Exhibit-TableFootnote',
        # End of chapter headings (Training data)
        'EOC-H1', 'EOC-H2',
        # EOC numbered lists (Official WK Template + Training)
        'EOC-NL-FIRST', 'EOC-NL-MID', 'EOC-NL-LAST',
        'EOC_NL', 'EOC_NLF', 'EOC_NLLL',
        'EOC-NumberList1', 'EOC-NumberList1_first', 'EOC-NumberLis1t_first', 'EOC-NumberList1_last',
        # EOC bulleted lists (Training data)
        'EOC-BL-FIRST', 'EOC-BL-MID', 'EOC-BL-LAST',
        'EOC-BulletList1', 'EOC-BulletList1_first', 'EOC-BulletList1_last', 'EOC-BulletList2',
        # EOC lettered lists (Training data)
        'EOC-Lc-AlphaList2', 'EOC-LL2-MID',
        # EOC dialogue (Training data)
        'EOC-Dialogue', 'EOC-UL-FL1',
        # EOC text/other (Training data)
        'EOC-Para-FL', 'EOC-ParaFirstLine-Ind', 'EOC-EQ-ONLY',
        'EOC_REF',
        # Glossary (Official WK Template)
        'GLOS-UL-FIRST', 'GLOS-UL-MID',
        'GLOS-NL-FIRST', 'GLOS-NL-MID',
        'GLOS-BL-FIRST', 'GLOS-BL-MID',
        # Index (Official WK Template)
        'IDX-TXT', 'IDX-ALPHA', 'IDX-1', 'IDX-2', 'IDX-3',
        # Appendix (Official WK Template + Training)
        'APX', 'APXN', 'APXT', 'APXST', 'APXAU',
        'APXH1', 'APXH2', 'APXH3',
        'APX-TXT', 'APX-TXT-FLUSH', 'APX-REF-N',
        # TOC elements (Official WK Template)
        'TOC-FM', 'TOC-UN', 'TOC-UT', 'TOC-SN', 'TOC-ST',
        'TOC-CN', 'TOC-CT', 'TOC-CAU',
        'TOC-H1', 'TOC-H2', 'TOC-BM-FIRST', 'TOC-BM',
        # Backmatter (Official WK Template)
        'BM-TTL',
        # Markers
        'PMI', 'ParaFirstLine-Ind',
    ],
    
    # EXERCISE: Exercise/workbook content (Official WK Template)
    'EXERCISE': [
        # Headers
        'EXER-H1', 'EXER-TTL', 'EXER-DIR',
        # True/False
        'EXER-TF-NL-FIRST', 'EXER-TF-NL-MID',
        # Multiple Choice
        'EXER-MC-NL-FIRST', 'EXER-MC-NL-MID', 'EXER-MC-NL2-FIRST', 'EXER-MC-NL2-MID',
        # Matching
        'EXER-M-NL-FIRST', 'EXER-M-NL-MID',
        # Fill Blank
        'EXER-FB-NL-FIRST', 'EXER-FB-NL-MID',
        # Short Answer
        'EXER-SA-NL-FIRST', 'EXER-SA-NL-MID',
        # Abbreviations
        'EXER-AB-NL-FIRST', 'EXER-AB-NL-MID',
        # Word Parts
        'EXER-WP-NL-FIRST', 'EXER-WP-NL-MID', 'EXER-WP-L',
        # Word Build
        'EXER-WB-NL-FIRST', 'EXER-WB-NL-MID',
        # Spelling
        'EXER-SP-NL-FIRST', 'EXER-SP-NL-MID', 'EXER-SP-NL2-FIRST', 'EXER-SP-NL2-MID',
        # Define Term
        'EXER-DT-NL-FIRST', 'EXER-DT-NL-MID',
        # Analyze
        'EXER-AT-NL-FIRST', 'EXER-AT-NL-MID', 'EXER-AT-T2',
        # Case Study
        'EXER-CS-AU', 'EXER-CS-T', 'EXER-CS-T2', 'EXER-CS-NL-FIRST', 'EXER-CS-NL-MID',
        # Other
        'EXER-L-UL',
        'PMI',
    ],
    
    # BODY has access to all styles (no restriction)
    'BODY': None,  # None means all styles allowed
}


def validate_style_for_zone(style: str, zone: str) -> bool:
    """
    Check if a style is valid for a given zone.
    
    Args:
        style: The style tag to validate
        zone: The context zone (FRONT_MATTER, BODY, TABLE, BOX_*, BACK_MATTER, EXERCISE)
    
    Returns:
        True if style is valid for zone, False otherwise
    """
    valid_styles = ZONE_VALID_STYLES.get(zone)
    
    # BODY has no restrictions
    if valid_styles is None:
        return True
    
    # Check exact match
    if style in valid_styles:
        return True
    
    # Check if style starts with any valid prefix (for variants like BL-MID0, TXT1, etc.)
    # Only for styles that might have numeric suffixes
    base_style = style.rstrip('0123456789')
    if base_style != style and base_style in valid_styles:
        return True
    
    return False


def get_valid_styles_for_zone(zone: str) -> list:
    """
    Get list of valid styles for a zone.
    
    Args:
        zone: The context zone
    
    Returns:
        List of valid style names, or empty list if no restrictions
    """
    styles = ZONE_VALID_STYLES.get(zone)
    return styles if styles else []


def get_zone_style_summary(zone: str) -> str:
    """
    Get a human-readable summary of valid styles for a zone.
    
    Args:
        zone: The context zone
    
    Returns:
        String summary of valid style prefixes
    """
    styles = ZONE_VALID_STYLES.get(zone)
    
    if styles is None:
        return "Full style range (no restrictions)"
    
    # Group by prefix for readability
    prefixes = set()
    for s in styles:
        if '-' in s:
            prefixes.add(s.split('-')[0] + '-*')
        else:
            prefixes.add(s)
    
    # Limit to most common prefixes
    common = sorted(prefixes)[:15]
    if len(prefixes) > 15:
        return ', '.join(common) + f' (+{len(prefixes)-15} more)'
    return ', '.join(common)


class DocumentIngestion:
    """
    Extract paragraphs and structure from DOCX files.
    Detects document context zones for better classification.
    
    Zones: FRONT_MATTER, BODY, BACK_MATTER, TABLE, BOX_*
    """
    
    def __init__(self, max_text_length: int = 200):
        """
        Initialize the document ingestion module.
        
        Args:
            max_text_length: Max characters for truncated display
        """
        self.max_text_length = max_text_length
        self._current_box_type = None  # Track current box context
    
    def _detect_box_start(self, text: str) -> Optional[str]:
        """
        Check if text starts a box and return box type.
        
        Returns:
            Box type (e.g., 'note', 'clinical pearl') or None
        """
        text_lower = text.lower().strip()
        
        for pattern in BOX_START_PATTERNS:
            if re.match(pattern, text_lower, re.IGNORECASE):
                # Extract box type from pattern
                for box_type in BOX_TYPE_MAPPING.keys():
                    if box_type.replace(' ', r'\s*') in pattern or box_type in text_lower:
                        return box_type
                # Default to 'box' if we can't determine type
                return 'box'
        return None
    
    def _detect_box_end(self, text: str) -> bool:
        """Check if text ends a box."""
        text_lower = text.lower().strip()
        
        for pattern in BOX_END_PATTERNS:
            if re.match(pattern, text_lower, re.IGNORECASE):
                return True
        return False
    
    def _get_box_zone(self, box_type: str) -> str:
        """Get the zone identifier for a box type."""
        if box_type is None:
            return 'BODY'
        
        prefix = BOX_TYPE_MAPPING.get(box_type, 'NBX')
        return f'BOX_{prefix}'
    
    def _detect_zone(self, text: str, current_zone: str, is_table: bool = False) -> tuple[str, Optional[str]]:
        """
        Detect which context zone this paragraph belongs to.
        
        Zone flow: METADATA → FRONT_MATTER → BODY → BACK_MATTER
        
        - METADATA: Pure metadata (book title, ORCID, etc.) - before chapter opener
        - FRONT_MATTER: Chapter opener through objectives (CN, CT, CAU, OBJ-*)
        - BODY: Main content starting at first H1
        - BACK_MATTER: References, index, etc.
        
        Args:
            text: Paragraph text
            current_zone: Current zone state
            is_table: Whether this is table content
            
        Returns:
            Tuple of (zone_identifier, box_type or None)
        """
        text_lower = text.lower().strip()
        
        # Table content always gets TABLE zone (unless in a box within table)
        if is_table:
            # Check if we're in a box context
            if self._current_box_type:
                # Check for box end
                if self._detect_box_end(text):
                    box_type = self._current_box_type
                    self._current_box_type = None
                    return (self._get_box_zone(box_type), box_type)
                return (self._get_box_zone(self._current_box_type), self._current_box_type)
            
            # Check for box start within table
            box_type = self._detect_box_start(text)
            if box_type:
                self._current_box_type = box_type
                return (self._get_box_zone(box_type), box_type)
            
            return ('TABLE', None)
        
        # Check for box end first (if we're in a box)
        if self._current_box_type:
            if self._detect_box_end(text):
                box_type = self._current_box_type
                self._current_box_type = None
                return (self._get_box_zone(box_type), box_type)
            return (self._get_box_zone(self._current_box_type), self._current_box_type)
        
        # Check for box start
        box_type = self._detect_box_start(text)
        if box_type:
            self._current_box_type = box_type
            return (self._get_box_zone(box_type), box_type)
        
        # If already in BODY, check for back matter
        if current_zone == 'BODY':
            for pattern in BACK_MATTER_PATTERNS:
                if re.match(pattern, text_lower, re.IGNORECASE):
                    return ('BACK_MATTER', None)
            return ('BODY', None)
        
        # If in BACK_MATTER, stay there
        if current_zone == 'BACK_MATTER':
            return ('BACK_MATTER', None)
        
        # Check for BODY start (first H1 heading)
        for pattern in BODY_START_PATTERNS:
            if re.match(pattern, text_lower, re.IGNORECASE):
                return ('BODY', None)
        
        # Check if this is a metadata section (pure PMI)
        if current_zone == 'METADATA':
            # Check for chapter opener - transitions to FRONT_MATTER
            for pattern in CHAPTER_OPENER_PATTERNS:
                if re.match(pattern, text_lower, re.IGNORECASE):
                    return ('FRONT_MATTER', None)
            
            # Check for end of metadata
            if '</metadata>' in text_lower:
                return ('METADATA', None)  # Still metadata, next will transition
            
            # Check if still metadata
            for pattern in METADATA_PATTERNS:
                if re.search(pattern, text_lower, re.IGNORECASE):
                    return ('METADATA', None)
            
            # If no metadata pattern but not chapter opener, stay in metadata
            return ('METADATA', None)
        
        # Check for chapter opener patterns - start of FRONT_MATTER
        for pattern in CHAPTER_OPENER_PATTERNS:
            if re.match(pattern, text_lower, re.IGNORECASE):
                return ('FRONT_MATTER', None)
        
        # If in FRONT_MATTER, stay there until H1
        if current_zone == 'FRONT_MATTER':
            return ('FRONT_MATTER', None)
        
        # Default: stay in current zone
        return (current_zone, None)
    
    # Keep old method for backward compatibility
    def _detect_section(self, text: str, current_section: str, prev_texts: list) -> str:
        """Legacy method - wraps _detect_zone for backward compatibility."""
        zone, _ = self._detect_zone(text, current_section, is_table=False)
        # Convert BOX zones back to BODY for old interface
        if zone.startswith('BOX_'):
            return 'BODY'
        if zone == 'METADATA':
            return 'FRONT_MATTER'
        return zone
    
    def extract_paragraphs(self, docx_path: str | Path) -> list[dict]:
        """
        Extract all paragraphs from a DOCX file with context zone detection.
        
        Zone flow: METADATA → FRONT_MATTER → BODY → BACK_MATTER
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List of paragraph dictionaries with id, text, metadata including:
            - context_zone: METADATA, FRONT_MATTER, BODY, BACK_MATTER, TABLE, BOX_*
            - box_type: If in a box, the type (note, clinical pearl, etc.)
            - valid_styles: List of styles valid for this zone (None = all allowed)
        """
        doc = Document(docx_path)
        paragraphs = []
        para_id = 1
        self._current_box_type = None  # Reset box tracking
        
        # Determine starting zone based on first paragraph
        all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Check what the document starts with
        if all_texts:
            first_text_lower = all_texts[0].lower()
            
            # Check for metadata patterns (pure PMI stuff)
            is_metadata_start = any(
                re.search(p, first_text_lower, re.IGNORECASE) 
                for p in METADATA_PATTERNS
            )
            
            # Check for chapter opener (CN/CT)
            is_chapter_start = any(
                re.match(p, first_text_lower, re.IGNORECASE) 
                for p in CHAPTER_OPENER_PATTERNS
            )
            
            # Check for H1/body start
            is_body_start = any(
                re.match(p, first_text_lower, re.IGNORECASE) 
                for p in BODY_START_PATTERNS
            )
            
            if is_metadata_start:
                current_zone = 'METADATA'
            elif is_chapter_start:
                current_zone = 'FRONT_MATTER'
            elif is_body_start:
                current_zone = 'BODY'
            else:
                current_zone = 'FRONT_MATTER'  # Default to front matter
        else:
            current_zone = 'FRONT_MATTER'
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Skip completely empty paragraphs
            if not text:
                continue
            
            # Detect zone (includes box tracking)
            current_zone, box_type = self._detect_zone(text, current_zone, is_table=False)
            
            # Extract formatting metadata
            metadata = self._extract_formatting(para)
            
            # Add zone information
            metadata['context_zone'] = current_zone
            metadata['box_type'] = box_type
            metadata['valid_styles'] = ZONE_VALID_STYLES.get(current_zone)
            
            # Keep backward compatible field
            if current_zone.startswith('BOX_'):
                metadata['document_section'] = 'BODY'
            else:
                metadata['document_section'] = current_zone
            
            paragraphs.append({
                'id': para_id,
                'text': text,
                'text_truncated': self._truncate(text),
                'metadata': metadata
            })
            
            para_id += 1
        
        # Reset box tracking before processing tables
        self._current_box_type = None
        
        # Also extract table content with zone awareness
        table_paragraphs = self._extract_tables(doc, para_id, current_zone)
        paragraphs.extend(table_paragraphs)
        
        # Log zone breakdown
        zone_counts = {}
        for p in paragraphs:
            zone = p['metadata'].get('context_zone', 'UNKNOWN')
            zone_counts[zone] = zone_counts.get(zone, 0) + 1
        
        logger.info(f"Extracted {len(paragraphs)} paragraphs from {docx_path}")
        logger.info(f"Zone breakdown: {zone_counts}")
        
        return paragraphs
    
    def _extract_formatting(self, para) -> dict:
        """
        Extract formatting metadata from a paragraph.
        
        Args:
            para: python-docx Paragraph object
            
        Returns:
            Dictionary of formatting properties
        """
        metadata = {
            'style_name': para.style.name if para.style else 'Normal',
            'alignment': None,
            'is_bold': False,
            'is_italic': False,
            'is_all_caps': False,
            'font_size': None,
            'has_numbering': False,
            'has_bullet': False,
            'indent_level': 0,
        }
        
        # Check alignment
        if para.alignment:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.CENTER: 'center',
                WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
            }
            metadata['alignment'] = alignment_map.get(para.alignment, 'left')
        
        # Check run formatting (first run)
        if para.runs:
            run = para.runs[0]
            metadata['is_bold'] = run.bold or False
            metadata['is_italic'] = run.italic or False
            if run.font.size:
                metadata['font_size'] = run.font.size.pt
            if run.font.all_caps:
                metadata['is_all_caps'] = True
        
        # Check for numbering/bullets in text (manual lists)
        text = para.text.strip()
        if re.match(r'^\d+\.?\s', text):
            metadata['has_numbering'] = True
        if re.match(r'^[•\-\*]\s', text):
            metadata['has_bullet'] = True
            
        # Check for XML numbering (automatic Word lists)
        # This catches lists that don't have bullets/numbers in the text property
        try:
            if hasattr(para, '_p') and para._p.pPr is not None and para._p.pPr.numPr is not None:
                # It has a numbering property -> it is a list
                # Try to guess type from style name
                style_name = para.style.name.lower() if para.style else ""
                
                if 'bullet' in style_name:
                    metadata['has_bullet'] = True
                elif 'number' in style_name:
                    metadata['has_numbering'] = True
                else:
                    # Ambiguous list - mark as generic list
                    # DO NOT default to numbering as it causes false positives (BL becoming NL)
                    metadata['has_xml_list'] = True
        except Exception:
            pass  # Fail gracefully if XML access fails
            
        # Check paragraph format for indentation
        if para.paragraph_format.left_indent:
            indent_inches = para.paragraph_format.left_indent.inches
            metadata['indent_level'] = int(indent_inches / 0.25)  # Estimate level
        
        return metadata
    
    def _extract_tables(self, doc, start_id: int, current_zone: str = 'BODY') -> list[dict]:
        """
        Extract text content from tables with zone awareness.
        
        Args:
            doc: python-docx Document object
            start_id: Starting paragraph ID
            current_zone: Current document zone context
            
        Returns:
            List of paragraph dictionaries from tables
        """
        paragraphs = []
        para_id = start_id
        
        for table_idx, table in enumerate(doc.tables):
            num_rows = len(table.rows)
            num_cols = len(table.columns) if table.rows else 0
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Process each paragraph in the cell (cells can have multiple paragraphs)
                    for para_idx, para in enumerate(cell.paragraphs):
                        text = para.text.strip()
                        if not text:
                            continue
                        
                        # Get actual style from cell paragraph
                        cell_style = para.style.name if para.style else None
                        
                        # Determine position context
                        is_header = row_idx == 0
                        is_first_col = cell_idx == 0
                        is_last_row = row_idx == num_rows - 1
                        
                        # Check for box markers within table cells
                        cell_zone, box_type = self._detect_zone(text, current_zone, is_table=True)
                        
                        # Infer appropriate table style based on position
                        inferred_style = self._infer_table_style(
                            cell_style, 
                            text, 
                            is_header, 
                            is_first_col,
                            is_last_row,
                            para_idx
                        )
                        
                        paragraphs.append({
                            'id': para_id,
                            'text': text,
                            'text_truncated': self._truncate(text),
                            'metadata': {
                                'style_name': cell_style or 'TableCell',
                                'inferred_style': inferred_style,
                                'is_table': True,
                                'table_index': table_idx,
                                'row_index': row_idx,
                                'cell_index': cell_idx,
                                'para_in_cell': para_idx,
                                'is_header_row': is_header,
                                'is_first_column': is_first_col,
                                'table_size': f"{num_rows}x{num_cols}",
                                'document_section': current_zone if not current_zone.startswith('BOX_') else 'BODY',
                                'context_zone': cell_zone,
                                'box_type': box_type,
                                'valid_styles': ZONE_VALID_STYLES.get(cell_zone),
                            }
                        })
                        para_id += 1
        
        return paragraphs
    
    def _infer_table_style(
        self, 
        cell_style: str, 
        text: str, 
        is_header: bool, 
        is_first_col: bool,
        is_last_row: bool,
        para_idx: int
    ) -> str:
        """
        Infer appropriate WK template table style based on cell position and content.
        """
        # If cell has a recognized style, use it
        if cell_style:
            style_upper = cell_style.upper()
            # Map common table styles
            if style_upper in ['T', 'TABLEBODY', 'GT']:
                if is_first_col:
                    return 'T4'  # Row header
                return 'T'  # Body cell
            elif style_upper in ['T2', 'TABLECOLUMNHEAD1', 'TABLEHEADER']:
                return 'T2'  # Column header
            elif 'TBL' in style_upper or 'BULLET' in style_upper:
                return 'TBL-MID'  # Table bullet list
            elif style_upper in ['TFN', 'TABLEFOOTNOTE']:
                return 'TFN'
            elif style_upper.startswith('UNT'):
                if is_header:
                    return 'T2'
                elif is_first_col:
                    return 'T4'
                return 'T'
        
        # Infer from position
        if is_header:
            return 'T2'  # Header row
        elif is_first_col:
            return 'T4'  # Row header (first column)
        
        # Check if text looks like a list item
        if text.startswith(('•', '-', '●', '○', '\t•', '\t-', '	')):
            return 'TBL-MID'
        
        # Default to body cell
        return 'T'
    
    def _truncate(self, text: str) -> str:
        """Truncate text for display purposes."""
        if len(text) <= self.max_text_length:
            return text
        return text[:self.max_text_length] + "..."
    
    def format_for_prompt(self, paragraphs: list[dict]) -> str:
        """
        Format paragraphs for the Gemini prompt.
        
        Args:
            paragraphs: List of paragraph dictionaries
            
        Returns:
            Formatted string for the prompt
        """
        lines = []
        for para in paragraphs:
            lines.append(f"[{para['id']}] {para['text_truncated']}")
        return "\n".join(lines)
    
    def get_document_stats(self, paragraphs: list[dict]) -> dict:
        """
        Get statistics about the extracted document.
        
        Args:
            paragraphs: List of paragraph dictionaries
            
        Returns:
            Dictionary of statistics
        """
        total_chars = sum(len(p['text']) for p in paragraphs)
        table_paras = sum(1 for p in paragraphs if p['metadata'].get('is_table'))
        numbered = sum(1 for p in paragraphs if p['metadata'].get('has_numbering'))
        bulleted = sum(1 for p in paragraphs if p['metadata'].get('has_bullet'))
        
        return {
            'total_paragraphs': len(paragraphs),
            'total_characters': total_chars,
            'estimated_tokens': total_chars // 4,  # Rough estimate
            'table_paragraphs': table_paras,
            'numbered_items': numbered,
            'bulleted_items': bulleted,
        }


def extract_document(docx_path: str | Path) -> tuple[list[dict], dict]:
    """
    Convenience function to extract a document.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Tuple of (paragraphs, stats)
    """
    ingestion = DocumentIngestion()
    paragraphs = ingestion.extract_paragraphs(docx_path)
    stats = ingestion.get_document_stats(paragraphs)
    return paragraphs, stats


if __name__ == "__main__":
    # Test with sample file
    import sys
    if len(sys.argv) > 1:
        path = sys.argv[1]
        paragraphs, stats = extract_document(path)
        print(f"Document Statistics: {stats}")
        print(f"\nFirst 10 paragraphs:")
        for p in paragraphs[:10]:
            print(f"  [{p['id']}] {p['text_truncated']}")
