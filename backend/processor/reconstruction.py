"""
STAGE 4: Document Reconstruction
- Style Applicator (applies Word paragraph styles)
- DOCX Builder
- Review Report Generator

Applies proper Word styles to paragraphs for professional output.
NO CONTENT MODIFICATION - only styles and formatting are applied.

Key implementation notes (critical for golden tests):
- Paragraph IDs must match ingestion order exactly:
  1) non-empty body paragraphs in document order
  2) then non-empty table cell paragraphs in document order (table->row->cell->paragraph)
- When processing tables, increment para_id per non-empty paragraph (NOT per cell).
"""

from __future__ import annotations

from datetime import datetime
import json
import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


# Style definitions with formatting properties - extracted from tagged documents
# Note: Text flow variants (TXT1, TXT2, H11, H12, etc.) inherit from base styles
STYLE_DEFINITIONS = {
    # Document Structure
    "CN": {"font_size": 24, "bold": True, "alignment": "center", "space_after": 6},
    "CT": {"font_size": 28, "bold": True, "alignment": "center", "space_after": 12},
    "CAU": {"font_size": 10, "italic": True, "alignment": "center", "space_after": 6},
    "PN": {"font_size": 10, "alignment": "center"},
    "PT": {"font_size": 14, "bold": True, "alignment": "center"},

    # Headings
    "H1": {"font_size": 16, "bold": True, "space_before": 12, "space_after": 6},
    "H2": {"font_size": 14, "bold": True, "space_before": 10, "space_after": 4},
    "H3": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "H4": {"font_size": 11, "bold": True, "space_before": 6, "space_after": 2},
    "H5": {"font_size": 11, "bold": True, "italic": True, "space_before": 6, "space_after": 2},
    "H6": {"font_size": 10, "bold": True, "space_before": 4, "space_after": 2},
    "SP-H1": {"font_size": 14, "bold": True, "caps": True, "space_before": 12, "space_after": 6},
    "EOC-H1": {"font_size": 14, "bold": True, "space_before": 12, "space_after": 6},
    # Heading text flow variants
    "H11": {"font_size": 16, "bold": True, "space_before": 12, "space_after": 6},
    "H12": {"font_size": 16, "bold": True, "space_before": 12, "space_after": 6},
    "H21": {"font_size": 14, "bold": True, "space_before": 10, "space_after": 4},

    # Reference Headings
    "REFH1": {"font_size": 14, "bold": True, "space_before": 12, "space_after": 6},
    "REFH2": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "REFH2a": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "Ref-H1": {"font_size": 14, "bold": True, "space_before": 12, "space_after": 6},
    "Ref-H2": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},

    # Body Text
    "TXT": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT-FLUSH": {"font_size": 11, "space_after": 6},
    "TXT-DC": {"font_size": 11, "space_after": 6},  # Drop cap would need special handling
    "TXT-AU": {"font_size": 10, "italic": True, "space_after": 6},
    "T": {"font_size": 10, "space_after": 2},  # Table cell body text
    # Body text flow variants
    "TXT1": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT2": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT3": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT4": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT-FLUSH1": {"font_size": 11, "space_after": 6},
    "TXT-FLUSH2": {"font_size": 11, "space_after": 6},
    "TXT-FLUSH4": {"font_size": 11, "space_after": 6},

    # Bulleted Lists
    "BL-FIRST": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "BL-MID": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "BL-LAST": {"font_size": 11, "left_indent": 0.5, "space_after": 6, "bullet": True},
    "UL-FIRST": {"font_size": 11, "left_indent": 0.75, "space_after": 2, "bullet": True},
    "UL-MID": {"font_size": 11, "left_indent": 0.75, "space_after": 2, "bullet": True},
    "UL-LAST": {"font_size": 11, "left_indent": 0.75, "space_after": 6, "bullet": True},

    # Numbered Lists
    "NL-FIRST": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "NL-MID": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "NL-LAST": {"font_size": 11, "left_indent": 0.5, "space_after": 6, "numbered": True},

    # End of Chapter Lists
    "EOC-NL-FIRST": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "EOC-NL-MID": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "EOC-NL-LAST": {"font_size": 11, "left_indent": 0.5, "space_after": 6, "numbered": True},
    "EOC-LL2-MID": {"font_size": 11, "left_indent": 0.75, "space_after": 2},

    # Tables - Titles
    "T1": {"font_size": 10, "bold": True, "space_before": 6, "space_after": 2},
    "T11": {"font_size": 10, "bold": True, "space_before": 6, "space_after": 2},
    "T12": {"font_size": 10, "bold": True, "space_before": 6, "space_after": 2},

    # Tables - Headers
    "T2": {"font_size": 10, "bold": True, "alignment": "center"},
    "T2-C": {"font_size": 10, "bold": True, "alignment": "center"},
    "T21": {"font_size": 10, "bold": True},  # Category/row headers
    "T22": {"font_size": 10, "bold": True, "alignment": "center"},  # Column headers
    "T23": {"font_size": 10, "bold": True, "alignment": "center"},  # Specific headers

    # Tables - Body Cells
    "T3": {"font_size": 10, "bold": True},  # Row header/subhead
    "T5": {"font_size": 10},  # Table body cell (data values)
    "T6": {"font_size": 10},  # Table body cell variant

    # Tables - Lists inside cells
    "TBL-FIRST": {"font_size": 10, "left_indent": 0.25, "space_after": 2, "bullet": True},
    "TBL-MID": {"font_size": 10, "left_indent": 0.25, "bullet": True},
    "TBL-MID0": {"font_size": 10, "left_indent": 0.25, "bullet": True},
    "TBL-LAST": {"font_size": 10, "left_indent": 0.25, "space_after": 4, "bullet": True},
    "TBL-LAST1": {"font_size": 10, "left_indent": 0.25, "space_after": 4, "bullet": True},
    "TBL2-MID": {"font_size": 10, "left_indent": 0.5, "bullet": True},
    "TBL3-MID": {"font_size": 10, "left_indent": 0.75, "bullet": True},
    "TBL4-MID": {"font_size": 10, "left_indent": 1.0, "bullet": True},
    "TUL-MID": {"font_size": 10, "left_indent": 0.25, "bullet": True},

    # Tables - Footnotes
    "TFN": {"font_size": 9, "italic": True, "space_before": 2},
    "TFN1": {"font_size": 9, "italic": True, "space_before": 2},
    "TSN": {"font_size": 9, "space_before": 2},

    # Figures
    "FIG-LEG": {"font_size": 10, "space_before": 6, "space_after": 6},
    "PMI": {"font_size": 9, "italic": True},

    # References
    "REF-N": {"font_size": 10, "hanging_indent": 0.5, "space_after": 2},
    "REF-N0": {"font_size": 10, "hanging_indent": 0.5, "space_after": 2},

    # Chapter Outline
    "COUT-1": {"font_size": 11, "left_indent": 0.25, "space_after": 2},
    "COUT-2": {"font_size": 11, "left_indent": 0.5, "space_after": 2},

    # Equations
    "EQ-ONLY": {"font_size": 11, "alignment": "center", "space_before": 6, "space_after": 6},
    "EQ-MID": {"font_size": 11, "alignment": "center", "space_after": 4},  # Multi-line equation continuation

    # Appendix Styles
    "APX-TYPE": {"font_size": 12, "bold": True, "caps": True, "space_before": 12, "space_after": 4},
    "APX-TTL": {"font_size": 14, "bold": True, "space_after": 6},
    "APX-H1": {"font_size": 16, "bold": True, "space_before": 12, "space_after": 6},
    "APX-H2": {"font_size": 14, "bold": True, "space_before": 10, "space_after": 4},
    "APX-H3": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "APX-TXT": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "APX-TXT-FLUSH": {"font_size": 11, "space_after": 6},
    "APX-CAU": {"font_size": 10, "italic": True, "space_after": 6},
    "APX-REF-N": {"font_size": 10, "hanging_indent": 0.5, "space_after": 2},
    "APX-REFH1": {"font_size": 14, "bold": True, "space_before": 12, "space_after": 6},

    # Box Content
    "NBX1-TTL": {"font_size": 12, "bold": True, "space_after": 4},
    "NBX1-TXT": {"font_size": 10, "first_line_indent": 0.25, "space_after": 4},
    "NBX1-TXT-FLUSH": {"font_size": 10, "space_after": 4},
    "NBX1-BL-FIRST": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "NBX1-BL-MID": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "NBX1-BL-LAST": {"font_size": 10, "left_indent": 0.5, "space_after": 4, "bullet": True},
    "NBX1-BL2-MID": {"font_size": 10, "left_indent": 0.75, "space_after": 2, "bullet": True},
    "NBX1-NL-FIRST": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "NBX1-NL-MID": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "NBX1-NL-LAST": {"font_size": 10, "left_indent": 0.5, "space_after": 4, "numbered": True},
    "NBX1-DIA-FIRST": {"font_size": 10, "left_indent": 0.5, "space_after": 2},
    "NBX1-DIA-MID": {"font_size": 10, "left_indent": 0.5, "space_after": 2},
    "NBX1-DIA-LAST": {"font_size": 10, "left_indent": 0.5, "space_after": 4},
    "NBX1-UNT": {"font_size": 10, "space_after": 4},
    "NBX1-UNT-T2": {"font_size": 10, "bold": True, "space_after": 2},
    "NBX1-SRC": {"font_size": 9, "italic": True, "space_after": 4},
    "BX1-TXT-FIRST": {"font_size": 10, "space_after": 4},

    # Case Studies
    "CS-H1": {"font_size": 14, "bold": True, "space_before": 10, "space_after": 4},
    "CS-TTL": {"font_size": 12, "bold": True, "space_after": 4},
    "CS-TXT": {"font_size": 10, "first_line_indent": 0.25, "space_after": 4},
    "CS-TXT-FLUSH": {"font_size": 10, "space_after": 4},
    "CS-QUES-TXT": {"font_size": 10, "bold": True, "space_after": 2},
    "CS-ANS-TXT": {"font_size": 10, "space_after": 4},

    # Learning Objectives
    "OBJ1": {"font_size": 12, "bold": True, "space_after": 4},
    "OBJ-TXT": {"font_size": 10, "space_after": 4},
    "OBJ-BL-FIRST": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "OBJ-BL-MID": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "OBJ-BL-LAST": {"font_size": 10, "left_indent": 0.5, "space_after": 4, "bullet": True},

    # Special/Miscellaneous
    "SUMHD": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "EXT-ONLY": {"font_size": 10, "italic": True, "left_indent": 0.5, "space_after": 6},
}


class DocumentReconstructor:
    """
    Apply style tags to a DOCX and generate outputs.
    IMPORTANT: This module must not change content; only apply styles/markers.
    """

    def __init__(self, output_dir: str | Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True, parents=True)

    def _get_or_create_style(self, doc: Document, style_name: str) -> None:
        """
        Create paragraph style if missing and apply formatting from STYLE_DEFINITIONS (when available).
        """
        styles = doc.styles

        # Already exists
        try:
            _ = styles[style_name]
            return
        except KeyError:
            pass

        style_def = STYLE_DEFINITIONS.get(style_name, {})

        try:
            new_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            new_style.base_style = styles["Normal"]

            font = new_style.font
            font.size = Pt(style_def.get("font_size", 11))
            font.bold = style_def.get("bold", False)
            font.italic = style_def.get("italic", False)
            font.all_caps = style_def.get("caps", False)

            if "color" in style_def:
                font.color.rgb = RGBColor.from_string(style_def["color"])

            pf = new_style.paragraph_format

            if "alignment" in style_def:
                alignment_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                pf.alignment = alignment_map.get(style_def["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

            if "space_before" in style_def:
                pf.space_before = Pt(style_def["space_before"])
            if "space_after" in style_def:
                pf.space_after = Pt(style_def["space_after"])

            if "first_line_indent" in style_def:
                pf.first_line_indent = Inches(style_def["first_line_indent"])
            if "left_indent" in style_def:
                pf.left_indent = Inches(style_def["left_indent"])
            if "hanging_indent" in style_def:
                pf.first_line_indent = Inches(-style_def["hanging_indent"])
                pf.left_indent = Inches(style_def["hanging_indent"])

            logger.debug("Created style: %s", style_name)
        except Exception as e:
            logger.warning("Could not create style '%s': %s", style_name, e)

    def ensure_paragraph_style(self, doc: Document, style_name: str):
        """
        Return a paragraph style object for `style_name`.
        If it doesn't exist, create it (with formatting if available).
        """
        self._get_or_create_style(doc, style_name)
        return doc.styles[style_name]

    @staticmethod
    def _iter_nonempty_body_paragraphs(doc: Document):
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                yield para

    @staticmethod
    def _iter_nonempty_table_paragraphs(doc: Document):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text and para.text.strip():
                            yield para

    def apply_styles(self, source_path: str | Path, classifications: list[dict], output_name: Optional[str] = None) -> Path:
        """
        Apply classification tags as Word paragraph styles to BOTH body and table paragraphs.
        MUST mirror ingestion paragraph order.
        """
        source_path = Path(source_path)
        doc = Document(source_path)

        clf_lookup = {int(c["id"]): c for c in classifications}

        # Pre-create all styles referenced by classifications
        for tag in {c["tag"] for c in classifications}:
            self._get_or_create_style(doc, tag)

        # Content integrity: compare non-empty paragraph counts (body + table)
        original_body_count = sum(1 for _ in self._iter_nonempty_body_paragraphs(doc))
        original_table_para_count = sum(1 for _ in self._iter_nonempty_table_paragraphs(doc))
        logger.info("Content integrity (before): %s body paras, %s table paras", original_body_count, original_table_para_count)

        para_id = 1

        # Body paragraphs
        for para in self._iter_nonempty_body_paragraphs(doc):
            clf = clf_lookup.get(para_id)
            if clf:
                tag = clf["tag"]
                conf = int(clf.get("confidence", 85))
                para.style = self.ensure_paragraph_style(doc, tag)
                if conf < 85:
                    self._highlight_for_review(para)
            para_id += 1

        # Table paragraphs (IMPORTANT: para_id increments per non-empty table paragraph)
        for para in self._iter_nonempty_table_paragraphs(doc):
            clf = clf_lookup.get(para_id)
            if clf:
                tag = clf["tag"]
                conf = int(clf.get("confidence", 85))
                para.style = self.ensure_paragraph_style(doc, tag)
                if conf < 85:
                    self._highlight_for_review(para)
            para_id += 1

        final_body_count = sum(1 for _ in self._iter_nonempty_body_paragraphs(doc))
        final_table_para_count = sum(1 for _ in self._iter_nonempty_table_paragraphs(doc))
        logger.info("Content integrity (after): %s body paras, %s table paras", final_body_count, final_table_para_count)

        if final_body_count != original_body_count or final_table_para_count != original_table_para_count:
            raise ValueError(
                f"Content integrity check failed: body {original_body_count}->{final_body_count}, "
                f"table {original_table_para_count}->{final_table_para_count}"
            )

        if output_name is None:
            output_name = f"{source_path.stem}_tagged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        output_path = self.output_dir / output_name
        doc.save(output_path)
        logger.info("Saved styled document to %s", output_path)
        return output_path

    @staticmethod
    def _highlight_for_review(para):
        """
        Yellow highlight for low-confidence items.
        """
        try:
            for run in para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        except Exception:
            pass

    def apply_tags_with_markers(
        self, source_path: str | Path, classifications: list[dict], output_name: Optional[str] = None
    ) -> Path:
        """
        Apply tags as <TAG> markers at the start of EACH non-empty paragraph (body + table).
        Also applies paragraph style equal to the tag name.
        """
        source_path = Path(source_path)
        doc = Document(source_path)

        clf_lookup = {int(c["id"]): c for c in classifications}

        # Pre-create styles referenced by classifications (so "para.style = tag" always works)
        for tag in {c["tag"] for c in classifications}:
            self._get_or_create_style(doc, tag)

        def _has_any_marker(text: str) -> bool:
            return text.lstrip().startswith("<")

        para_id = 1

        # Body paragraphs
        for para in self._iter_nonempty_body_paragraphs(doc):
            clf = clf_lookup.get(para_id)
            if clf:
                tag = clf["tag"]
                para.style = self.ensure_paragraph_style(doc, tag)

                if not _has_any_marker(para.text):
                    # Prefix marker without removing existing text
                    if para.runs:
                        para.runs[0].text = f"<{tag}> " + para.runs[0].text
                    else:
                        para.text = f"<{tag}> {para.text}"
            para_id += 1

        # Table paragraphs (apply to EACH non-empty table paragraph)
        for para in self._iter_nonempty_table_paragraphs(doc):
            clf = clf_lookup.get(para_id)
            if clf:
                tag = clf["tag"]
                para.style = self.ensure_paragraph_style(doc, tag)

                if not _has_any_marker(para.text):
                    if para.runs:
                        para.runs[0].text = f"<{tag}> " + para.runs[0].text
                    else:
                        para.text = f"<{tag}> {para.text}"
            para_id += 1

        if output_name is None:
            output_name = f"{source_path.stem}_tagged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        output_path = self.output_dir / output_name
        doc.save(output_path)
        logger.info("Saved tagged document with markers to %s", output_path)
        return output_path

    def generate_review_report(self, document_name: str, filtered_results: dict, output_name: Optional[str] = None) -> Path:
        doc = Document()
        doc.add_heading("Pre-Editor Review Report", level=0)
        doc.add_paragraph(f"Source Document: {document_name}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()

        doc.add_heading("Summary", level=1)
        summary = filtered_results.get("summary", {})
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = "Table Grid"

        rows_data = [
            ("Total Paragraphs", str(summary.get("total_paragraphs", 0))),
            ("Auto-Applied", str(summary.get("auto_applied", 0))),
            ("Needs Review", str(summary.get("needs_review", 0))),
            ("Auto-Apply Rate", f"{summary.get('auto_apply_percentage', 0):.1f}%"),
        ]
        for i, (label, value) in enumerate(rows_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        needs_review = filtered_results.get("needs_review", [])
        if needs_review:
            doc.add_heading("Items Requiring Review", level=1)
            doc.add_paragraph(
                f"The following {len(needs_review)} items have confidence below 85% and require human review."
            )
            doc.add_paragraph()

            for item in needs_review:
                p = doc.add_paragraph()
                run = p.add_run(f"Paragraph {item['id']}")
                run.bold = True

                txt = item.get("original_text", "")
                doc.add_paragraph(f'Text: "{(txt[:100] + "...") if len(txt) > 100 else txt}"')
                doc.add_paragraph(f"Suggested Tag: {item.get('tag', 'Unknown')} (Confidence: {item.get('confidence', 0)}%)")

                if item.get("reasoning"):
                    doc.add_paragraph(f"Reasoning: {item['reasoning']}")
                if item.get("alternatives"):
                    doc.add_paragraph(f"Alternative Tags: {', '.join(item['alternatives'])}")
                doc.add_paragraph()
        else:
            doc.add_heading("All Items Auto-Applied", level=1)
            doc.add_paragraph("All paragraphs were classified with high confidence (≥85%). No manual review required.")

        if output_name is None:
            output_name = f"review_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        output_path = self.output_dir / output_name
        doc.save(output_path)
        logger.info("Saved review report to %s", output_path)
        return output_path

    def generate_json_output(
        self,
        document_name: str,
        classifications: list[dict],
        filtered_results: dict,
        output_name: Optional[str] = None,
    ) -> Path:
        output_data = {
            "document_name": document_name,
            "processed_at": datetime.now().isoformat(),
            "summary": filtered_results.get("summary", {}),
            "classifications": classifications,
            "flagged_items": filtered_results.get("needs_review", []),
        }
        if output_name is None:
            output_name = f"classification_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

        output_path = self.output_dir / output_name
        output_path.write_text(json.dumps(output_data, indent=2), encoding="utf-8")
        logger.info("Saved JSON results to %s", output_path)
        return output_path

    def generate_html_report(
        self,
        document_name: str,
        classifications: list[dict],
        filtered_results: dict,
        output_name: Optional[str] = None,
    ) -> Path:
        from .html_report import generate_html_report

        if output_name is None:
            output_name = f"classification_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"

        output_path = self.output_dir / output_name
        generate_html_report(document_name, classifications, filtered_results, output_path)
        logger.info("Saved HTML report to %s", output_path)
        return output_path


def reconstruct_document(
    source_path: str | Path,
    classifications: list[dict],
    filtered_results: dict,
    output_dir: str | Path,
    use_markers: bool = False,
    output_base: str | None = None,
) -> dict:
    """
    Convenience function to generate all outputs.
    """
    source_path = Path(source_path)
    reconstructor = DocumentReconstructor(output_dir)

    if output_base is None:
        output_base = f"{source_path.stem}_processed"

    tagged_name = f"{output_base}.docx"
    report_name = f"{output_base}_review.docx"
    json_name = f"{output_base}_results.json"
    html_name = f"{output_base}_report.html"

    if use_markers:
        tagged_path = reconstructor.apply_tags_with_markers(source_path, classifications, tagged_name)
    else:
        tagged_path = reconstructor.apply_styles(source_path, classifications, tagged_name)

    report_path = reconstructor.generate_review_report(source_path.name, filtered_results, report_name)
    json_path = reconstructor.generate_json_output(source_path.name, classifications, filtered_results, json_name)
    html_path = reconstructor.generate_html_report(source_path.name, classifications, filtered_results, html_name)

    return {
        "tagged_document": tagged_path.name,
        "review_report": report_path.name,
        "json_results": json_path.name,
        "html_report": html_path.name,
    }
