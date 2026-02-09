import os
import io
import sys
import json
import zipfile
from pathlib import Path
from collections import Counter, defaultdict

from docx import Document

repo_root = Path(__file__).resolve().parents[1]
backend_path = repo_root / "backend"
sys.path.insert(0, str(backend_path))
sys.path.insert(0, str(repo_root))
print(f"[eval_accuracy] backend path: {backend_path}")

from processor.pipeline import process_document
from processor.ingestion import extract_document
from app.services.style_normalizer import normalize_style
from app.services.reference_zone import detect_reference_zone

RAW_ZIP = repo_root / "Org files.zip"
TAG_ZIP = repo_root / "manual tagged files.zip"
OUTPUT_ROOT = repo_root / "outputs" / "eval"
ZONE_BASED_DIR = repo_root / "test zone-based"
GROUND_TRUTH_FILE = backend_path / "data" / "ground_truth.jsonl"

DEV_BASES = {
    "Acharya9781975261764-ch002",
    "Goroll9781975212643-ch227",
    "Lynn9781975252823-ch004",
}
HELDOUT_BASES = {
    "Howley9781975221171-ch071",
    "Bittner9781975243012-sec6.9",
}


def load_ground_truth():
    """
    Load ground truth dataset from ground_truth.jsonl.
    Returns a dict: {doc_id: [entries sorted by para_index]}
    """
    if not GROUND_TRUTH_FILE.exists():
        raise FileNotFoundError(f"Ground truth file not found: {GROUND_TRUTH_FILE}")

    ground_truth = defaultdict(list)
    total_entries = 0
    unmapped_count = 0

    with open(GROUND_TRUTH_FILE, 'r', encoding='utf-8') as f:
        for line in f:
            if not line.strip():
                continue
            entry = json.loads(line)
            doc_id = entry["doc_id"]
            ground_truth[doc_id].append(entry)
            total_entries += 1
            if entry.get("canonical_gold_tag") == "UNMAPPED":
                unmapped_count += 1

    # Sort each doc's entries by para_index
    for doc_id in ground_truth:
        ground_truth[doc_id].sort(key=lambda e: e["para_index"])

    mapped_count = total_entries - unmapped_count
    mapped_pct = (mapped_count / total_entries * 100) if total_entries else 0

    print(f"\n=== Ground Truth Dataset ===")
    print(f"Gold source: manual tagged files.zip")
    print(f"Total entries: {total_entries}")
    print(f"Mapped entries: {mapped_count} ({mapped_pct:.1f}%)")
    print(f"UNMAPPED entries: {unmapped_count} ({100-mapped_pct:.1f}%)")
    print(f"Documents: {len(ground_truth)}")
    print()

    return dict(ground_truth)


def _extract_style_sequence(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    styles: list[str] = []

    for para in doc.paragraphs:
        if para.text and para.text.strip():
            styles.append(para.style.name if para.style else "Normal")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text and para.text.strip():
                        styles.append(para.style.name if para.style else "Normal")

    return styles


def _collect_pairs():
    pairs = []

    # Collect from zips
    with zipfile.ZipFile(RAW_ZIP) as raw_zip, zipfile.ZipFile(TAG_ZIP) as tag_zip:
        raw_names = [n for n in raw_zip.namelist() if n.lower().endswith("_org.docx")]
        tag_candidates = [n for n in tag_zip.namelist() if n.lower().endswith(".docx") and ("_tag" in n.lower())]
        tag_by_base = {}
        for n in tag_candidates:
            base = Path(n).name
            base_key = base.replace("_tagged.docx", "").replace("_tag.docx", "")
            tag_by_base[base_key] = n

        for raw_name in raw_names:
            base_key = Path(raw_name).name.replace("_org.docx", "")
            if base_key in tag_by_base:
                pairs.append((base_key, raw_name, tag_by_base[base_key]))

    # Collect from test zone-based folder (preferred for fixed split)
    if ZONE_BASED_DIR.exists():
        files = list(ZONE_BASED_DIR.glob("*.docx"))
        by_base = {}
        for f in files:
            name = f.name
            base = (
                name.replace(" org.docx", "")
                .replace("_org.docx", "")
                .replace(" tag.docx", "")
                .replace("_tag.docx", "")
                .replace("_tagged.docx", "")
            )
            by_base.setdefault(base, {})["raw" if "org" in name.lower() else "tag"] = f

        for base, items in by_base.items():
            raw = items.get("raw")
            tag = items.get("tag")
            if raw and tag:
                pairs.append((base, raw, tag))

    return pairs


def _select_pairs(pairs):
    selected = []
    for base, raw_name, tag_name in pairs:
        if base in DEV_BASES or base in HELDOUT_BASES:
            selected.append((base, raw_name, tag_name))
    return selected


def _find_doc_id_match(base_name, ground_truth):
    """
    Find the matching doc_id in ground truth for a given base name.
    Handles variations like:
    - base_name: "Acharya9781975261764-ch002"
    - doc_id: "Acharya9781975261764-ch002_tag"
    """
    # Try exact match first
    if base_name in ground_truth:
        return base_name

    # Try with _tag suffix
    if f"{base_name}_tag" in ground_truth:
        return f"{base_name}_tag"

    # Try finding partial match
    for doc_id in ground_truth:
        if base_name in doc_id or doc_id in base_name:
            return doc_id

    return None


def _evaluate_predictions_against_gold(base_name, gold_entries, prediction_blocks):
    """
    Evaluate predictions against gold standard entries aligned by para_index.

    Args:
        base_name: document base name
        gold_entries: list of ground truth entries sorted by para_index
        prediction_blocks: list of blocks from extract_document with classifications

    Returns:
        dict with metrics: total, matches, accuracy, mismatches, coverage_gaps, zone_violations, unmapped_count
    """
    # Index predictions by para_index
    # Note: prediction_blocks come from extract_document which processes the raw file
    # The para_index should match the ground truth para_index if both are based on the same document structure
    predictions_by_idx = {}

    # Try to get para_index from metadata first
    has_para_index = any(block.get("metadata", {}).get("para_index") is not None for block in prediction_blocks)

    if has_para_index:
        # Use para_index from metadata
        for block in prediction_blocks:
            para_idx = block.get("metadata", {}).get("para_index")
            if para_idx is not None:
                predicted_tag = normalize_style(block.get("metadata", {}).get("style_name", "TXT"))
                predictions_by_idx[para_idx] = predicted_tag
    else:
        # Fall back to using block order as para_index
        # Filter out empty blocks
        non_empty_blocks = [b for b in prediction_blocks if b.get("text", "").strip()]
        for idx, block in enumerate(non_empty_blocks):
            predicted_tag = normalize_style(block.get("metadata", {}).get("style_name", "TXT"))
            predictions_by_idx[idx] = predicted_tag

    # Evaluate
    total = 0
    matches = 0
    unmapped_count = 0
    mismatches = Counter()
    coverage_gaps = []  # (gold_tag, para_index)
    zone_violations = []  # {zone, gold_tag, predicted_tag, para_index}

    # Zone constraints from classifier.py
    ZONE_CONSTRAINTS = {
        'TABLE': {
            'T', 'T1', 'T2', 'T2-C', 'T3', 'T4', 'T5', 'TD',
            'TH1', 'TH2', 'TH3', 'TH4', 'TH5', 'TH6',
            'T-DIR', 'T1-C', 'TN', 'TSN', 'T-ALT', 'TSN-C',
        },
        'BOX': {'BX-TXT', 'BX-H', 'BX-BL-FIRST', 'BX-BL-MID', 'BX-BL-LAST'},
        'BOX1': {'BX1-TXT', 'BX1-H', 'BX1-BL-FIRST', 'BX1-BL-MID', 'BX1-BL-LAST'},
        'BOX2': {'BX2-TXT', 'BX2-H', 'BX2-BL-FIRST', 'BX2-BL-MID', 'BX2-BL-LAST'},
        'BOX3': {'BX3-TXT', 'BX3-H', 'BX3-BL-FIRST', 'BX3-BL-MID', 'BX3-BL-LAST'},
        'BOX4': {'BX4-TXT', 'BX4-H', 'BX4-BL-FIRST', 'BX4-BL-MID', 'BX4-BL-LAST'},
        'BACK_MATTER': {
            'REF-N', 'REF-U', 'REFH1', 'UNFIG-LEG', 'TSN', 'BM-TTL',
            'FN', 'FN-LBL', 'IND-H1', 'IND-H2', 'IND-BL-FIRST', 'IND-BL-MID', 'IND-BL-LAST',
        },
    }

    for entry in gold_entries:
        para_idx = entry["para_index"]
        gold_tag = entry["canonical_gold_tag"]
        zone = entry.get("zone", "body").upper()

        # Skip UNMAPPED entries in accuracy calculation
        if gold_tag == "UNMAPPED":
            unmapped_count += 1
            continue

        total += 1

        # Check if we have a prediction for this para_index
        if para_idx not in predictions_by_idx:
            coverage_gaps.append((gold_tag, para_idx))
            continue

        predicted_tag = predictions_by_idx[para_idx]

        # Compare
        if gold_tag == predicted_tag:
            matches += 1
        else:
            mismatches[(gold_tag, predicted_tag)] += 1

        # Check zone violation
        zone_valid_tags = ZONE_CONSTRAINTS.get(zone)
        if zone_valid_tags and predicted_tag not in zone_valid_tags:
            zone_violations.append({
                "zone": zone,
                "gold_tag": gold_tag,
                "predicted_tag": predicted_tag,
                "para_index": para_idx
            })

    accuracy = (matches / total * 100) if total > 0 else 0

    return {
        "total": total,
        "matches": matches,
        "accuracy": accuracy,
        "mismatches": mismatches,
        "coverage_gaps": coverage_gaps,
        "zone_violations": zone_violations,
        "unmapped_count": unmapped_count,
    }


def _print_report(
    base,
    metrics,
    ref_zone_count,
    ref_trigger,
    ref_examples,
    ref_start_idx,
    ref_start_ratio,
):
    """
    Print evaluation report for a single document.

    Args:
        base: document base name
        metrics: dict with keys: total, matches, accuracy, mismatches, coverage_gaps, zone_violations, unmapped_count
        ref_zone_count: number of reference zone blocks
        ref_trigger: reference zone detection trigger
        ref_examples: list of reference zone example snippets
        ref_start_idx: reference zone start index
        ref_start_ratio: reference zone start ratio
    """
    total = metrics["total"]
    matches = metrics["matches"]
    accuracy = metrics["accuracy"]
    mismatches = metrics["mismatches"]
    coverage_gaps = metrics["coverage_gaps"]
    zone_violations = metrics["zone_violations"]
    unmapped_count = metrics["unmapped_count"]

    print("\n=== Eval Report ===")
    print("Doc:", base)
    print(f"Total gold entries: {total}")
    print(f"Matched predictions: {matches}")
    print(f"Accuracy: {accuracy:.2f}%")
    print(f"UNMAPPED gold entries: {unmapped_count}")
    print(f"Coverage gaps (no prediction): {len(coverage_gaps)}")
    print(f"Zone violations: {len(zone_violations)}")
    print(f"Reference zone blocks: {ref_zone_count} (trigger={ref_trigger})")
    if ref_start_idx is not None:
        print(f"Reference zone start_idx: {ref_start_idx} (start_ratio={ref_start_ratio:.2f})")
    else:
        print("Reference zone start_idx: None")
    if ref_examples:
        print("Reference zone examples:")
        for ref_id, snippet in ref_examples:
            print(f"  [{ref_id}] {snippet}")

    print("\nTop 20 mismatches (gold → predicted):")
    for (gold, pred), count in mismatches.most_common(20):
        print(f"  {gold} → {pred}: {count}")

    if coverage_gaps:
        print("\nTop 10 coverage gaps (gold tag, no prediction):")
        gap_counter = Counter([gold_tag for gold_tag, _ in coverage_gaps])
        for gold_tag, count in gap_counter.most_common(10):
            print(f"  {gold_tag}: {count}")

    if zone_violations:
        print("\nTop 10 zone violations (predicted tag not valid for zone):")
        violation_counter = Counter([(v["zone"], v["predicted_tag"]) for v in zone_violations])
        for (zone, pred_tag), count in violation_counter.most_common(10):
            print(f"  {zone}: {pred_tag} ({count}x)")


def main():
    if not RAW_ZIP.exists() or not TAG_ZIP.exists():
        raise SystemExit("Missing input zips.")

    # Load ground truth dataset
    ground_truth = load_ground_truth()

    pairs = _collect_pairs()
    selected = _select_pairs(pairs)
    # De-duplicate by base name and keep fixed split only once
    seen = set()
    unique = []
    for base, raw_name, tag_name in selected:
        if base in seen:
            continue
        seen.add(base)
        unique.append((base, raw_name, tag_name))
    selected = unique

    print("Selected docs:", [b for b, _, _ in selected])

    if len(selected) < 5:
        raise SystemExit("Not enough selected pairs found in zips.")

    dev_mismatches = Counter()
    dev_coverage_gaps = []
    dev_zone_violations = []
    dev_matches = 0
    dev_total = 0

    held_mismatches = Counter()
    held_coverage_gaps = []
    held_zone_violations = []
    held_matches = 0
    held_total = 0

    with zipfile.ZipFile(RAW_ZIP) as raw_zip, zipfile.ZipFile(TAG_ZIP) as tag_zip:
        for base, raw_name, tag_name in selected:
            try:
                # Check if we have ground truth for this doc
                doc_id = _find_doc_id_match(base, ground_truth)
                if not doc_id:
                    print(f"\n[SKIP] {base}: No ground truth data found")
                    continue

                gold_entries = ground_truth[doc_id]
                print(f"\n[PROCESSING] {base} (gold entries: {len(gold_entries)})")

                temp_dir = OUTPUT_ROOT / "_tmp" / base
                temp_dir.mkdir(parents=True, exist_ok=True)

                if isinstance(raw_name, Path):
                    raw_path = raw_name
                else:
                    raw_path = Path(raw_zip.extract(raw_name, temp_dir))

                if isinstance(tag_name, Path):
                    tagged_path = tag_name
                else:
                    tagged_path = Path(tag_zip.extract(tag_name, temp_dir))

                output_dir = OUTPUT_ROOT / base
                (output_dir / "processed").mkdir(parents=True, exist_ok=True)
                (output_dir / "review").mkdir(parents=True, exist_ok=True)
                (output_dir / "json").mkdir(parents=True, exist_ok=True)

                result = process_document(
                    input_path=str(raw_path),
                    output_folder=str(output_dir),
                    use_markers=True,
                    apply_repair=True,
                )

                output_path = Path(result["output_path"]).parent / "output.docx"
                Path(result["output_path"]).replace(output_path)

                blocks, _ = extract_document(raw_path)
                ref_ret = detect_reference_zone(blocks)
                if isinstance(ref_ret, tuple):
                    ref_ids = ref_ret[0]
                    ref_trigger = ref_ret[1] if len(ref_ret) > 1 else "none"
                    ref_start_idx = ref_ret[2] if len(ref_ret) > 2 else None
                else:
                    ref_ids = ref_ret
                    ref_trigger = "none"
                    ref_start_idx = None
                ref_zone_count = len(ref_ids)
                ref_start_ratio = (ref_start_idx / len(blocks)) if ref_start_idx is not None and blocks else 0
                ref_examples = []
                for b in blocks:
                    if b.get("id") in ref_ids:
                        text = b.get("text", "")
                        ref_examples.append((b.get("id"), text[:60]))
                    if len(ref_examples) >= 5:
                        break

                # Get prediction blocks (output of pipeline with classifications)
                prediction_blocks, _ = extract_document(output_path)

                # Evaluate predictions against ground truth
                metrics = _evaluate_predictions_against_gold(base, gold_entries, prediction_blocks)

                if base.startswith("Goroll"):
                    out_ref_ret = detect_reference_zone(prediction_blocks)
                    if isinstance(out_ref_ret, tuple):
                        output_ref_ids = out_ref_ret[0]
                    else:
                        output_ref_ids = out_ref_ret
                    ul_bl_count = 0
                    for b in prediction_blocks:
                        if b.get("id") in output_ref_ids:
                            style_name = normalize_style(b.get("metadata", {}).get("style_name", ""))
                            if style_name.startswith("UL-") or style_name.startswith("BL-"):
                                ul_bl_count += 1
                    print(f"Reference zone UL/BL count after repair (Goroll): {ul_bl_count}")

                _print_report(
                    base,
                    metrics,
                    ref_zone_count,
                    ref_trigger,
                    ref_examples,
                    ref_start_idx,
                    ref_start_ratio,
                )

                if base in DEV_BASES:
                    dev_mismatches.update(metrics["mismatches"])
                    dev_coverage_gaps.extend(metrics["coverage_gaps"])
                    dev_zone_violations.extend(metrics["zone_violations"])
                    dev_matches += metrics["matches"]
                    dev_total += metrics["total"]
                elif base in HELDOUT_BASES:
                    held_mismatches.update(metrics["mismatches"])
                    held_coverage_gaps.extend(metrics["coverage_gaps"])
                    held_zone_violations.extend(metrics["zone_violations"])
                    held_matches += metrics["matches"]
                    held_total += metrics["total"]
            except Exception as exc:
                print(f"[ERROR] {base}: {exc}")
                continue

    dev_acc = (dev_matches / dev_total) * 100 if dev_total else 0
    held_acc = (held_matches / held_total) * 100 if held_total else 0

    print("\n" + "=" * 60)
    print("=== OVERALL SUMMARY ===")
    print("=" * 60)

    print("\n--- DEV SET ---")
    print(f"Total gold entries: {dev_total}")
    print(f"Matched predictions: {dev_matches}")
    print(f"Accuracy: {dev_acc:.2f}%")
    print(f"Coverage gaps: {len(dev_coverage_gaps)}")
    print(f"Zone violations: {len(dev_zone_violations)}")

    print("\nDEV Top 30 mismatches (gold → predicted):")
    for (gold, pred), count in dev_mismatches.most_common(30):
        print(f"  {gold} → {pred}: {count}")

    if dev_coverage_gaps:
        print("\nDEV Top 20 coverage gaps (gold tag, no prediction):")
        gap_counter = Counter([gold_tag for gold_tag, _ in dev_coverage_gaps])
        for gold_tag, count in gap_counter.most_common(20):
            print(f"  {gold_tag}: {count}")

    if dev_zone_violations:
        print("\nDEV Top 20 zone violations:")
        violation_counter = Counter([(v["zone"], v["predicted_tag"]) for v in dev_zone_violations])
        for (zone, pred_tag), count in violation_counter.most_common(20):
            print(f"  {zone}: {pred_tag} ({count}x)")

    print("\n" + "-" * 60)
    print("--- HELDOUT SET ---")
    print(f"Total gold entries: {held_total}")
    print(f"Matched predictions: {held_matches}")
    print(f"Accuracy: {held_acc:.2f}%")
    print(f"Coverage gaps: {len(held_coverage_gaps)}")
    print(f"Zone violations: {len(held_zone_violations)}")

    print("\nHELDOUT Top 30 mismatches (gold → predicted):")
    for (gold, pred), count in held_mismatches.most_common(30):
        print(f"  {gold} → {pred}: {count}")

    if held_coverage_gaps:
        print("\nHELDOUT Top 20 coverage gaps (gold tag, no prediction):")
        gap_counter = Counter([gold_tag for gold_tag, _ in held_coverage_gaps])
        for gold_tag, count in gap_counter.most_common(20):
            print(f"  {gold_tag}: {count}")

    if held_zone_violations:
        print("\nHELDOUT Top 20 zone violations:")
        violation_counter = Counter([(v["zone"], v["predicted_tag"]) for v in held_zone_violations])
        for (zone, pred_tag), count in violation_counter.most_common(20):
            print(f"  {zone}: {pred_tag} ({count}x)")

    print("\n" + "=" * 60)


if __name__ == "__main__":
    main()
